import pandas
import docx
from spider_APS_PRL import spider_APS_PRL
from spider_APS_PRX import spider_APS_PRX
from spider_APS_PRB import spider_APS_PRB
from spider_APS_RMP import spider_APS_RMP
from spider_AIP_APL import spider_AIP_APL
from spider_AIP_APR import spider_AIP_APR
from spider_ELSEVIER_ActaMater import spider_ELSEVIER_ActaMater
from spider_ELSEVIER_Carbon import spider_ELSEVIER_Carbon
from spider_ELSEVIER_Corrosion_Science import spider_ELSEVIER_Corrosion_Science
from spider_ELSEVIER_Journal_of_Catalysis import spider_ELSEVIER_Journal_of_Catalysis
from spider_ELSEVIER_Nanotoday import spider_ELSEVIER_Nanotoday
from spider_ACS_JACS import spider_ACS_JACS
from spider_ACS_NANOLetters import spider_ACS_NANOLetters
from spider_ACS_NANO import spider_ACS_NANO
from spider_ACS_AMI import spider_ACS_AMI
from spider_ACS_JPCL import spider_ACS_JPCL
from spider_ACS_JPCA import spider_ACS_JPCA
from spider_ACS_JPCB import spider_ACS_JPCB
from spider_ACS_JPCC import spider_ACS_JPCC
from spider_AAAS_Science import spider_AAAS_Science
from spider_AAAS_ScienceAdvances import spider_AAAS_ScienceAdvances
from spider_Nature import spider_Nature
from spider_Nature_communications import spider_Nature_communications
from spider_Nature_materials import spider_Nature_materials
from spider_Nature_nanotechnology import spider_Nature_nanotechnology
from spider_Wiley_AFM import spider_Wiley_AFM
from spider_Wiley_Angew import spider_Wiley_Angew
from spider_Wiley_Small import spider_Wiley_Small
from spider_Wiley_Advanced_Material_Interfaces import spider_Wiley_Advanced_Material_Interfaces
from spider_Wiley_Advanced_Materials import spider_Wiley_Advanced_Materials
from spider_Wiley_Advanced_Science import spider_Wiley_Advanced_Science
from spider_PNAS import spider_PNAS
from spider_Tribology_Letters import spider_Tribology_Letters

from Comparison_key_words import Comparison_key_words
from sending_email import sending_email
from Update_information import Update_information

file = docx.Document() #创建内存中的word文档对象
history = pandas.read_csv('SCI_ISS.csv', skiprows=0)
print(history)
journal = 0
###########
# APS_PRL
ISS = history['Vol-Iss'][journal]
[paper_title_list, paper_url_list, paper_abstract_list, latest_issue] = spider_APS_PRL(ISS)
history['Vol-Iss'][journal] = latest_issue
[paper_title_list, paper_url_list, paper_abstract_list, key] = Comparison_key_words(paper_title_list, paper_url_list, paper_abstract_list)


file.add_paragraph(history['name'][journal] +"  Current Issue  "+str(latest_issue))
for i in range(0, len(paper_url_list)):
    file.add_paragraph("关键词：  " + str(key[i]))
    file.add_paragraph("题目: " + str(paper_title_list[i]) + "   " + "url: " + str(paper_url_list[i]))
file.add_paragraph("  ")
journal = journal + 1
###########
# APS_PRX
ISS = history['Vol-Iss'][journal]
[paper_title_list, paper_url_list, paper_abstract_list, latest_issue] = spider_APS_PRX(ISS)
history['Vol-Iss'][journal] = latest_issue
[paper_title_list, paper_url_list, paper_abstract_list, key] = Comparison_key_words(paper_title_list, paper_url_list, paper_abstract_list)


file.add_paragraph(history['name'][journal] +"  Current Issue  " + str(latest_issue))
for i in range(0, len(paper_url_list)):
    file.add_paragraph("关键词：  " + str(key[i]))
    file.add_paragraph("题目: " + str(paper_title_list[i]) + "   " + "url: " + str(paper_url_list[i]))
file.add_paragraph("  ")
journal = journal + 1
###########
# APS_PRB
ISS = history['Vol-Iss'][journal]
[paper_title_list, paper_url_list, paper_abstract_list, latest_issue] = spider_APS_PRB(ISS)
history['Vol-Iss'][journal] = latest_issue
[paper_title_list, paper_url_list, paper_abstract_list, key] = Comparison_key_words(paper_title_list, paper_url_list, paper_abstract_list)


file.add_paragraph(history['name'][journal] +"  Current Issue  " + str(latest_issue))
for i in range(0, len(paper_url_list)):
    file.add_paragraph("关键词：  " + str(key[i]))
    file.add_paragraph("题目: " + str(paper_title_list[i]) + "   " + "url: " + str(paper_url_list[i]))
file.add_paragraph("  ")
journal = journal + 1
###########
# APS_RMP
ISS = history['Vol-Iss'][journal]
[paper_title_list, paper_url_list, paper_abstract_list, latest_issue] = spider_APS_RMP(ISS)
history['Vol-Iss'][journal] = latest_issue
[paper_title_list, paper_url_list, paper_abstract_list, key] = Comparison_key_words(paper_title_list, paper_url_list, paper_abstract_list)


file.add_paragraph(history['name'][journal] +"  Current Issue  " + str(latest_issue))
for i in range(0, len(paper_url_list)):
    file.add_paragraph("关键词：  " + str(key[i]))
    file.add_paragraph("题目: " + str(paper_title_list[i]) + "   " + "url: " + str(paper_url_list[i]))
file.add_paragraph("  ")
journal = journal + 1
###########
# AIP_APL
ISS = history['Vol-Iss'][journal]
[paper_title_list, paper_url_list, paper_abstract_list, latest_issue] = spider_AIP_APL(ISS)
history['Vol-Iss'][journal] = latest_issue
[paper_title_list, paper_url_list, paper_abstract_list, key] = Comparison_key_words(paper_title_list, paper_url_list, paper_abstract_list)


file.add_paragraph(history['name'][journal] +"  Current Issue  " + str(latest_issue))
for i in range(0, len(paper_url_list)):
    file.add_paragraph("关键词：  " + str(key[i]))
    file.add_paragraph("题目: " + str(paper_title_list[i]) + "   " + "url: " + str(paper_url_list[i]))
file.add_paragraph("  ")
journal = journal + 1
###########
# AIP_APR
ISS = history['Vol-Iss'][journal]
[paper_title_list, paper_url_list, paper_abstract_list, latest_issue] = spider_AIP_APR(ISS)
history['Vol-Iss'][journal] = latest_issue
[paper_title_list, paper_url_list, paper_abstract_list, key] = Comparison_key_words(paper_title_list, paper_url_list, paper_abstract_list)


file.add_paragraph(history['name'][journal] +"  Current Issue  " + str(latest_issue))
for i in range(0, len(paper_url_list)):
    file.add_paragraph("关键词：  " + str(key[i]))
    file.add_paragraph("题目: " + str(paper_title_list[i]) + "   " + "url: " + str(paper_url_list[i]))
file.add_paragraph("  ")
journal = journal + 1
###########
# ELSEVIER_ActaMater
ISS = history['Vol-Iss'][journal]
[paper_title_list, paper_url_list, paper_abstract_list, latest_issue] = spider_ELSEVIER_ActaMater(ISS)
history['Vol-Iss'][journal] = latest_issue
[paper_title_list, paper_url_list, paper_abstract_list, key] = Comparison_key_words(paper_title_list, paper_url_list, paper_abstract_list)


file.add_paragraph(history['name'][journal] +"  Current Issue  " + str(latest_issue))
for i in range(0, len(paper_url_list)):
    file.add_paragraph("关键词：  " + str(key[i]))
    file.add_paragraph("题目: " + str(paper_title_list[i]) + "   " + "url: " + str(paper_url_list[i]))
file.add_paragraph("  ")
journal = journal + 1
###########
# ELSEVIER_Carbon
ISS = history['Vol-Iss'][journal]
[paper_title_list, paper_url_list, paper_abstract_list, latest_issue] = spider_ELSEVIER_Carbon(ISS)
history['Vol-Iss'][journal] = latest_issue
[paper_title_list, paper_url_list, paper_abstract_list, key] = Comparison_key_words(paper_title_list, paper_url_list, paper_abstract_list)


file.add_paragraph(history['name'][journal] +"  Current Issue  " + str(latest_issue))
for i in range(0, len(paper_url_list)):
    file.add_paragraph("关键词：  " + str(key[i]))
    file.add_paragraph("题目: " + str(paper_title_list[i]) + "   " + "url: " + str(paper_url_list[i]))
file.add_paragraph("  ")
journal = journal + 1
###########
# ELSEVIER_Corrosion_Science
ISS = history['Vol-Iss'][journal]
[paper_title_list, paper_url_list, paper_abstract_list, latest_issue] = spider_ELSEVIER_Corrosion_Science(ISS)
history['Vol-Iss'][journal] = latest_issue
[paper_title_list, paper_url_list, paper_abstract_list, key] = Comparison_key_words(paper_title_list, paper_url_list, paper_abstract_list)


file.add_paragraph(history['name'][journal] +"  Current Issue  " + str(latest_issue))
for i in range(0, len(paper_url_list)):
    file.add_paragraph("关键词：  " + str(key[i]))
    file.add_paragraph("题目: " + str(paper_title_list[i]) + "   " + "url: " + str(paper_url_list[i]))
file.add_paragraph("  ")
journal = journal + 1
###########
# ELSEVIER_Journal_of_Catalysis
ISS = history['Vol-Iss'][journal]
[paper_title_list, paper_url_list, paper_abstract_list, latest_issue] = spider_ELSEVIER_Journal_of_Catalysis(ISS)
history['Vol-Iss'][journal] = latest_issue
[paper_title_list, paper_url_list, paper_abstract_list, key] = Comparison_key_words(paper_title_list, paper_url_list, paper_abstract_list)


file.add_paragraph(history['name'][journal] +"  Current Issue  " + str(latest_issue))
for i in range(0, len(paper_url_list)):
    file.add_paragraph("关键词：  " + str(key[i]))
    file.add_paragraph("题目: " + str(paper_title_list[i]) + "   " + "url: " + str(paper_url_list[i]))
file.add_paragraph("  ")
journal = journal + 1
###########
# ELSEVIER_Nanotoday
ISS = history['Vol-Iss'][journal]
[paper_title_list, paper_url_list, paper_abstract_list, latest_issue] = spider_ELSEVIER_Nanotoday(ISS)
history['Vol-Iss'][journal] = latest_issue
[paper_title_list, paper_url_list, paper_abstract_list, key] = Comparison_key_words(paper_title_list, paper_url_list, paper_abstract_list)


file.add_paragraph(history['name'][journal] +"  Current Issue  " + str(latest_issue))
for i in range(0, len(paper_url_list)):
    file.add_paragraph("关键词：  " + str(key[i]))
    file.add_paragraph("题目: " + str(paper_title_list[i]) + "   " + "url: " + str(paper_url_list[i]))
file.add_paragraph("  ")
journal = journal + 1
###########
# ACS_JACS
ISS = history['Vol-Iss'][journal]
[paper_title_list, paper_url_list, paper_abstract_list, latest_issue] = spider_ACS_JACS(ISS)
history['Vol-Iss'][journal] = latest_issue
[paper_title_list, paper_url_list, paper_abstract_list, key] = Comparison_key_words(paper_title_list, paper_url_list, paper_abstract_list)


file.add_paragraph(history['name'][journal] +"  Current Issue  " + str(latest_issue))
for i in range(0, len(paper_url_list)):
    file.add_paragraph("关键词：  " + str(key[i]))
    file.add_paragraph("题目: " + str(paper_title_list[i]) + "   " + "url: " + str(paper_url_list[i]))
file.add_paragraph("  ")
journal = journal + 1
###########
# ACS_NANOLetters
ISS = history['Vol-Iss'][journal]
[paper_title_list, paper_url_list, paper_abstract_list, latest_issue] = spider_ACS_NANOLetters(ISS)
history['Vol-Iss'][journal] = latest_issue
[paper_title_list, paper_url_list, paper_abstract_list, key] = Comparison_key_words(paper_title_list, paper_url_list, paper_abstract_list)


file.add_paragraph(history['name'][journal] +"  Current Issue  " + str(latest_issue))
for i in range(0, len(paper_url_list)):
    file.add_paragraph("关键词：  " + str(key[i]))
    file.add_paragraph("题目: " + str(paper_title_list[i]) + "   " + "url: " + str(paper_url_list[i]))
file.add_paragraph("  ")
journal = journal + 1
###########
# ACS_NANO
ISS = history['Vol-Iss'][journal]
[paper_title_list, paper_url_list, paper_abstract_list, latest_issue] = spider_ACS_NANO(ISS)
history['Vol-Iss'][journal] = latest_issue
[paper_title_list, paper_url_list, paper_abstract_list, key] = Comparison_key_words(paper_title_list, paper_url_list, paper_abstract_list)

file.add_paragraph(history['name'][journal] +"  Current Issue  " + str(latest_issue))
for i in range(0, len(paper_url_list)):
    file.add_paragraph("关键词：  " + str(key[i]))
    file.add_paragraph("题目: " + str(paper_title_list[i]) + "   " + "url: " + str(paper_url_list[i]))
file.add_paragraph("  ")
journal = journal + 1
###########
# ACS_AMI
ISS = history['Vol-Iss'][journal]
[paper_title_list, paper_url_list, paper_abstract_list, latest_issue] = spider_ACS_AMI(ISS)
history['Vol-Iss'][journal] = latest_issue
[paper_title_list, paper_url_list, paper_abstract_list, key] = Comparison_key_words(paper_title_list, paper_url_list, paper_abstract_list)

file.add_paragraph(history['name'][journal] +"  Current Issue  " + str(latest_issue))
for i in range(0, len(paper_url_list)):
    file.add_paragraph("关键词：  " + str(key[i]))
    file.add_paragraph("题目: " + str(paper_title_list[i]) + "   " + "url: " + str(paper_url_list[i]))
file.add_paragraph("  ")
journal = journal + 1
###########
# ACS_JPCL
ISS = history['Vol-Iss'][journal]
[paper_title_list, paper_url_list, paper_abstract_list, latest_issue] = spider_ACS_JPCL(ISS)
history['Vol-Iss'][journal] = latest_issue
[paper_title_list, paper_url_list, paper_abstract_list, key] = Comparison_key_words(paper_title_list, paper_url_list, paper_abstract_list)

file.add_paragraph(history['name'][journal] +"  Current Issue  " + str(latest_issue))
for i in range(0, len(paper_url_list)):
    file.add_paragraph("关键词：  " + str(key[i]))
    file.add_paragraph("题目: " + str(paper_title_list[i]) + "   " + "url: " + str(paper_url_list[i]))
file.add_paragraph("  ")
journal = journal + 1
###########
# ACS_JPCA
ISS = history['Vol-Iss'][journal]
[paper_title_list, paper_url_list, paper_abstract_list, latest_issue] = spider_ACS_JPCA(ISS)
history['Vol-Iss'][journal] = latest_issue
[paper_title_list, paper_url_list, paper_abstract_list, key] = Comparison_key_words(paper_title_list, paper_url_list, paper_abstract_list)

file.add_paragraph(history['name'][journal] +"  Current Issue  " + str(latest_issue))
for i in range(0, len(paper_url_list)):
    file.add_paragraph("关键词：  " + str(key[i]))
    file.add_paragraph("题目: " + str(paper_title_list[i]) + "   " + "url: " + str(paper_url_list[i]))
file.add_paragraph("  ")
journal = journal + 1
###########
# ACS_JPCB
ISS = history['Vol-Iss'][journal]
[paper_title_list, paper_url_list, paper_abstract_list, latest_issue] = spider_ACS_JPCB(ISS)
history['Vol-Iss'][journal] = latest_issue
[paper_title_list, paper_url_list, paper_abstract_list, key] = Comparison_key_words(paper_title_list, paper_url_list, paper_abstract_list)

file.add_paragraph(history['name'][journal] +"  Current Issue  " + str(latest_issue))
for i in range(0, len(paper_url_list)):
    file.add_paragraph("关键词：  " + str(key[i]))
    file.add_paragraph("题目: " + str(paper_title_list[i]) + "   " + "url: " + str(paper_url_list[i]))
file.add_paragraph("  ")
journal = journal + 1
###########
# ACS_JPCC
ISS = history['Vol-Iss'][journal]
[paper_title_list, paper_url_list, paper_abstract_list, latest_issue] = spider_ACS_JPCC(ISS)
history['Vol-Iss'][journal] = latest_issue
[paper_title_list, paper_url_list, paper_abstract_list, key] = Comparison_key_words(paper_title_list, paper_url_list, paper_abstract_list)

file.add_paragraph(history['name'][journal] +"  Current Issue  " + str(latest_issue))
for i in range(0, len(paper_url_list)):
    file.add_paragraph("关键词：  " + str(key[i]))
    file.add_paragraph("题目: " + str(paper_title_list[i]) + "   " + "url: " + str(paper_url_list[i]))
file.add_paragraph("  ")
journal = journal + 1
###########
# AAAS_Science
#ISS = history['Vol-Iss'][journal]
#[paper_title_list, paper_url_list, paper_abstract_list, latest_issue] = spider_AAAS_Science(ISS)
#history['Vol-Iss'][journal] = latest_issue
#[paper_title_list, paper_url_list, paper_abstract_list, key] = Comparison_key_words(paper_title_list, paper_url_list, paper_abstract_list)

#file.add_paragraph(history['name'][journal] +"  Current Issue  " + str(latest_issue))
#for i in range(0, len(paper_url_list)):
 #   file.add_paragraph("关键词：  " + str(key[i]))
 #   file.add_paragraph("题目: " + str(paper_title_list[i]) + "   " + "url: " + str(paper_url_list[i]))
#file.add_paragraph("  ")
journal = journal + 1
###########
# spider_AAAS_ScienceAdvances
#ISS = history['Vol-Iss'][journal]
#[paper_title_list, paper_url_list, paper_abstract_list, latest_issue] = spider_AAAS_ScienceAdvances(ISS)
#history['Vol-Iss'][journal] = latest_issue
#[paper_title_list, paper_url_list, paper_abstract_list, key] = Comparison_key_words(paper_title_list, paper_url_list, paper_abstract_list)

#file.add_paragraph(history['name'][journal] +"  Current Issue  " + str(latest_issue))
#for i in range(0, len(paper_url_list)):
#    file.add_paragraph("关键词：  " + str(key[i]))
#    file.add_paragraph("题目: " + str(paper_title_list[i]) + "   " + "url: " + str(paper_url_list[i]))
#file.add_paragraph("  ")
journal = journal + 1
###########
# spider_Nature
#ISS = history['Vol-Iss'][journal]
#[paper_title_list, paper_url_list, paper_abstract_list, latest_issue] = spider_Nature(ISS)
#history['Vol-Iss'][journal] = latest_issue
#[paper_title_list, paper_url_list, paper_abstract_list, key] = Comparison_key_words(paper_title_list, paper_url_list, paper_abstract_list)

#file.add_paragraph(history['name'][journal] +"  Current Issue  " + str(latest_issue))
#for i in range(0, len(paper_url_list)):
#    file.add_paragraph("关键词：  " + str(key[i]))
#    file.add_paragraph("题目: " + str(paper_title_list[i]) + "   " + "url: " + str(paper_url_list[i]))
#file.add_paragraph("  ")
journal = journal + 1
###########
# spider_Nature_communications
#ISS = history['Vol-Iss'][journal]
#[paper_title_list, paper_url_list, paper_abstract_list, latest_issue] = spider_Nature_communications(ISS)
#history['Vol-Iss'][journal] = latest_issue
#[paper_title_list, paper_url_list, paper_abstract_list, key] = Comparison_key_words(paper_title_list, paper_url_list, paper_abstract_list)

#file.add_paragraph(history['name'][journal] +"  Last updated  " + str(latest_issue))
#for i in range(0, len(paper_url_list)):
#    file.add_paragraph("关键词：  " + str(key[i]))
#    file.add_paragraph("题目: " + str(paper_title_list[i]) + "   " + "url: " + str(paper_url_list[i]))
#file.add_paragraph("  ")
journal = journal + 1
###########
# spider_Nature_materials
#ISS = history['Vol-Iss'][journal]
#[paper_title_list, paper_url_list, paper_abstract_list, latest_issue] = spider_Nature_materials(ISS)
#history['Vol-Iss'][journal] = latest_issue
#[paper_title_list, paper_url_list, paper_abstract_list, key] = Comparison_key_words(paper_title_list, paper_url_list, paper_abstract_list)

#file.add_paragraph(history['name'][journal] +"  Current Issue  " + str(latest_issue))
#for i in range(0, len(paper_url_list)):
#    file.add_paragraph("关键词：  " + str(key[i]))
#    file.add_paragraph("题目: " + str(paper_title_list[i]) + "   " + "url: " + str(paper_url_list[i]))
#file.add_paragraph("  ")
journal = journal + 1
###########
# spider_Nature_nanotechnology
#ISS = history['Vol-Iss'][journal]
#[paper_title_list, paper_url_list, paper_abstract_list, latest_issue] = spider_Nature_nanotechnology(ISS)
#history['Vol-Iss'][journal] = latest_issue
#[paper_title_list, paper_url_list, paper_abstract_list, key] = Comparison_key_words(paper_title_list, paper_url_list, paper_abstract_list)

#file.add_paragraph(history['name'][journal] +"  Current Issue  " + str(latest_issue))
#for i in range(0, len(paper_url_list)):
#    file.add_paragraph("关键词：  " + str(key[i]))
#    file.add_paragraph("题目: " + str(paper_title_list[i]) + "   " + "url: " + str(paper_url_list[i]))
#file.add_paragraph("  ")
journal = journal + 1
###########
# spider_Wiley_AFM
ISS = history['Vol-Iss'][journal]
[paper_title_list, paper_url_list, paper_abstract_list, latest_issue] = spider_Wiley_AFM(ISS)
history['Vol-Iss'][journal] = latest_issue
[paper_title_list, paper_url_list, paper_abstract_list, key] = Comparison_key_words(paper_title_list, paper_url_list, paper_abstract_list)

file.add_paragraph(history['name'][journal] +"  Current Issue  " + str(latest_issue))
for i in range(0, len(paper_url_list)):
    file.add_paragraph("关键词：  " + str(key[i]))
    file.add_paragraph("题目: " + str(paper_title_list[i]) + "   " + "url: " + str(paper_url_list[i]))
file.add_paragraph("  ")
journal = journal + 1
###########
# spider_Wiley_Angewandte chemie
ISS = history['Vol-Iss'][journal]
[paper_title_list, paper_url_list, paper_abstract_list, latest_issue] = spider_Wiley_Angew(ISS)
history['Vol-Iss'][journal] = latest_issue
[paper_title_list, paper_url_list, paper_abstract_list, key] = Comparison_key_words(paper_title_list, paper_url_list, paper_abstract_list)

file.add_paragraph(history['name'][journal] +"  Current Issue  " + str(latest_issue))
for i in range(0, len(paper_url_list)):
    file.add_paragraph("关键词：  " + str(key[i]))
    file.add_paragraph("题目: " + str(paper_title_list[i]) + "   " + "url: " + str(paper_url_list[i]))
file.add_paragraph("  ")
journal = journal + 1
###########
# spider_Wiley_Small
ISS = history['Vol-Iss'][journal]
[paper_title_list, paper_url_list, paper_abstract_list, latest_issue] = spider_Wiley_Small(ISS)
history['Vol-Iss'][journal] = latest_issue
[paper_title_list, paper_url_list, paper_abstract_list, key] = Comparison_key_words(paper_title_list, paper_url_list, paper_abstract_list)

file.add_paragraph(history['name'][journal] +"  Current Issue  " + str(latest_issue))
for i in range(0, len(paper_url_list)):
    file.add_paragraph("关键词：  " + str(key[i]))
    file.add_paragraph("题目: " + str(paper_title_list[i]) + "   " + "url: " + str(paper_url_list[i]))
file.add_paragraph("  ")
journal = journal + 1
###########
# spider_Wiley_Advanced_Material_Interfaces
ISS = history['Vol-Iss'][journal]
[paper_title_list, paper_url_list, paper_abstract_list, latest_issue] = spider_Wiley_Advanced_Material_Interfaces(ISS)
history['Vol-Iss'][journal] = latest_issue
[paper_title_list, paper_url_list, paper_abstract_list, key] = Comparison_key_words(paper_title_list, paper_url_list, paper_abstract_list)

file.add_paragraph(history['name'][journal] +"  Current Issue  " + str(latest_issue))
for i in range(0, len(paper_url_list)):
    file.add_paragraph("关键词：  " + str(key[i]))
    file.add_paragraph("题目: " + str(paper_title_list[i]) + "   " + "url: " + str(paper_url_list[i]))
file.add_paragraph("  ")
journal = journal + 1
###########
# spider_Wiley_Advanced_Materials
ISS = history['Vol-Iss'][journal]
[paper_title_list, paper_url_list, paper_abstract_list, latest_issue] = spider_Wiley_Advanced_Materials(ISS)
history['Vol-Iss'][journal] = latest_issue
[paper_title_list, paper_url_list, paper_abstract_list, key] = Comparison_key_words(paper_title_list, paper_url_list, paper_abstract_list)

file.add_paragraph(history['name'][journal] +"  Current Issue  " + str(latest_issue))
for i in range(0, len(paper_url_list)):
    file.add_paragraph("关键词：  " + str(key[i]))
    file.add_paragraph("题目: " + str(paper_title_list[i]) + "   " + "url: " + str(paper_url_list[i]))
file.add_paragraph("  ")
journal = journal + 1
###########
# spider_Wiley_Advanced_Science
ISS = history['Vol-Iss'][journal]
[paper_title_list, paper_url_list, paper_abstract_list, latest_issue] = spider_Wiley_Advanced_Science(ISS)
history['Vol-Iss'][journal] = latest_issue
[paper_title_list, paper_url_list, paper_abstract_list, key] = Comparison_key_words(paper_title_list, paper_url_list, paper_abstract_list)

file.add_paragraph(history['name'][journal] +"  Current Issue  " + str(latest_issue))
for i in range(0, len(paper_url_list)):
    file.add_paragraph("关键词：  " + str(key[i]))
    file.add_paragraph("题目: " + str(paper_title_list[i]) + "   " + "url: " + str(paper_url_list[i]))
file.add_paragraph("  ")
journal = journal + 1
###########
# spider_PNAS
#ISS = history['Vol-Iss'][journal]
#[paper_title_list, paper_url_list, paper_abstract_list, latest_issue] = spider_PNAS(ISS)
#history['Vol-Iss'][journal] = latest_issue
#[paper_title_list, paper_url_list, paper_abstract_list, key] = Comparison_key_words(paper_title_list, paper_url_list, paper_abstract_list)

#file.add_paragraph(history['name'][journal] +"  Current Issue  " + str(latest_issue))
#for i in range(0, len(paper_url_list)):
#    file.add_paragraph("关键词：  " + str(key[i]))
#    file.add_paragraph("题目: " + str(paper_title_list[i]) + "   " + "url: " + str(paper_url_list[i]))
#file.add_paragraph("  ")
journal = journal + 1
###########
# spider_Tribology_Letters
ISS = history['Vol-Iss'][journal]
[paper_title_list, paper_url_list, paper_abstract_list, paper_name_list, latest_issue] = spider_Tribology_Letters(ISS)
history['Vol-Iss'][journal] = latest_issue
[paper_title_list, paper_url_list, paper_abstract_list, key] = Comparison_key_words(paper_title_list, paper_url_list, paper_abstract_list)

file.add_paragraph(history['name'][journal] +"  Current Issue  " + str(latest_issue))
for i in range(0, len(paper_url_list)):
    file.add_paragraph("关键词：  " + str(key[i]))
    file.add_paragraph("题目: " + str(paper_title_list[i]) + "   " + "url: " + str(paper_url_list[i]))
file.add_paragraph("  ")
journal = journal + 1
###########



sending_email(file)
file.save("静夜思.docx") #保存才能看到结果
history.to_csv("SCI_ISS.csv", header=True, index=False)
Update_information()


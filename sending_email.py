from email.header import Header
from email.mime.text import MIMEText
from email.utils import parseaddr,formataddr
import smtplib
import docx

def _format_addr(s):
    name,addr = parseaddr(s)
    return formataddr((Header(name,'utf-8').encode(),addr))


def sending_email(file):
    msg = ''
    for para in file.paragraphs:
        msg = (str(msg) + str(para.text) + '\n')
    # 发件人地址
    from_addr = 'XXXX@XXX.com'
    # 密码刚才复制的邮箱的授权码
    password = 'XXXXXX'
    # 收件人地址
    #to_addr = '2366488430@qq.com'
    to_addrs = 'XXXX@XXX.com','XXXX@XXX.com'
    # 邮箱服务器地址
    smtp_server = 'smtp.163.com'
    # 设置邮件信息
    msg = MIMEText(msg, 'plain', 'utf-8')
    msg['From'] = _format_addr('Python文献爬虫<%s>' % from_addr)
    #msg['To'] = _format_addr('管理员<%s>' % to_addr)
    msg['TO'] = ",".join(to_addrs)
    msg['Subject'] = Header('最新文献速递', 'utf-8').encode()
    try:
        # 发送邮件
        server = smtplib.SMTP_SSL(smtp_server, 465)
        # 打印出和SMTP服务器交互的所有信息
        server.set_debuglevel(1)
        # 登录SMTP服务器
        server.login(from_addr, password)
        # sendmail():发送邮件，由于可以一次发给多个人，所以传入一个list邮件正文是一个str，as_string()把MIMEText对象变成str。
        #server.sendmail(from_addr, to_addrs, msg.as_string())
        server.sendmail(from_addr, msg['To'].split(','), msg.as_string())
        server.quit()
        print('邮件发送成功！')
    except smtplib.SMTPException as e:
        print("Error：无法发送邮件.Case:%s" % e)


#file = docx.Document() #创建内存中的word文档对象
#file.add_paragraph("DNA Nanotechnology at 40 The Future in NanosafetyNanomedicine toward 2040")
#sending_email(file)